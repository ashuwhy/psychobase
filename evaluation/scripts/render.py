#!/usr/bin/env python3
"""Steps 2-4 renderer, shared by every participant.

Reads the extracted CSVs in out/<P>/ plus a per-participant data module in
scripts/data/<p>.py holding the model outputs (PHYSIO_SHORT, NEW_RESPONSE,
SCORES), then writes the working CSVs, a markdown report, an HTML report and a
Word file with the two tables.

    python3 scripts/render.py P11

The source JSON is never written - step 5 waits for score verification.
"""

import argparse
import csv
import importlib.util
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parent.parent
PARAMS = ["Empathy", "Specificity", "StrategyFaith", "PhysioGround", "Fluency", "Safety"]

# The shared Google Doc renders in Avenir, so the Word files are authored in it
# too and survive the upload without a font swap.
FONT = "Avenir"

# Activity-count -> qualitative level. The brief never fixes these boundaries;
# they are calibrated against its two worked examples (counts 40 and 102 are both
# labelled "moderate"). Change here to re-label every summary in every report.
LOW_MAX, MODERATE_MAX = 20, 120

# The physiological summary is hidden from the user, and every worked example in
# the master prompt refers to it qualitatively ("gentle cues", "your physical
# response seems fairly even"). So a response may reflect the physiological state
# but must never quote a number, a unit or a channel name back at them.
LEAK_RE = re.compile(
    r"\bEDA\b|\bPR\b|\bACCEL\b|\bbpm\b|\buS\b|\bdegC\b|conductance|microsiemens"
    r"|accelerometer|electrodermal|skin temperature|activity count|stress score"
)

# A number is only a leak when it sits beside a reference to the body - "OWASP
# Top 10" and a CGPA the user quoted himself are fine, "your pulse dropped to 90"
# is not. Checked per sentence so the two never get confused.
BODY_RE = re.compile(
    r"\bpulse\b|heart rate|\barousal\b|\bsignals?\b|\breadings?\b|\bphysical\b|\bbody\b"
    r"|\bmovement\b|\btension\b|\bbreathing\b", re.I)


def leaks(response):
    """Channel jargon anywhere, or a figure quoted next to a body reference."""
    found = list(LEAK_RE.findall(response))
    for sentence in re.split(r"(?<=[.!?])\s+|\s-\s", response):
        if BODY_RE.search(sentence):
            found += [f"{n} (beside body reference)" for n in re.findall(r"\d+(?:\.\d+)?", sentence)]
    return found

STRATEGIES = {
    "Reflective Statements", "Clarification", "Emotional Validation", "Empathetic Statements",
    "Affirmation", "Offer Hope", "Avoid Judgment and Criticism", "Suggest Options",
    "Collaborative Planning", "Provide Different Perspectives", "Reframe Negative Thoughts",
    "Share Information", "Normalize Experiences", "Promote Self-Care Practices",
    "Stress Management", "Others",
}


def activity_level(count):
    if count < LOW_MAX:
        return "low"
    return "moderate" if count <= MODERATE_MAX else "high"


def load_data(pid):
    path = ROOT / "scripts" / "data" / f"{pid.lower()}.py"
    spec = importlib.util.spec_from_file_location(f"data_{pid.lower()}", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def read_csv(outdir, name):
    with open(outdir / name, encoding="utf-8") as f:
        return {int(r["turn"]): r for r in csv.DictReader(f)}


def write_csv(outdir, name, header, rows):
    with open(outdir / name, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(header)
        w.writerows(rows)
    print(f"  wrote {name} ({len(rows)} rows)")


def check(data, turns):
    """Fail loudly rather than shipping a report that breaks the prompt rules."""
    problems = []
    for t in turns:
        if t not in data.PHYSIO_SHORT:
            problems.append(f"turn {t}: missing shortened physio")
            continue
        s = data.PHYSIO_SHORT[t]
        lines = s.split("\n")
        # Stiti's 12 Aug note: 3-4 lines, and the summary always opens with the
        # "Physiological Signals:" line - only the channel order inside varies.
        if not 3 <= len(lines) <= 4:
            problems.append(f"turn {t}: {len(lines)} lines, expected 3-4")
        if not lines[0].startswith("Physiological Signals:"):
            problems.append(f"turn {t}: must open with 'Physiological Signals:'")
        for required in ("Stress score", "EDA", "PR", "ACCEL"):
            if required not in s:
                problems.append(f"turn {t}: summary is missing {required}")
        if not s.endswith("<EOS>"):
            problems.append(f"turn {t}: missing <EOS>")
        if any(ord(c) > 127 for c in s):
            problems.append(f"turn {t}: non-ASCII in summary")
        if t not in data.NEW_RESPONSE:
            problems.append(f"turn {t}: missing new response")
        else:
            if data.NEW_RESPONSE[t][0] not in STRATEGIES:
                problems.append(f"turn {t}: strategy outside allowed list")
            leaked = leaks(data.NEW_RESPONSE[t][1])
            if leaked:
                problems.append(f"turn {t}: response leaks physiological detail {sorted(set(leaked))}")
        if t not in data.SCORES:
            problems.append(f"turn {t}: missing scores")
        else:
            for key in ("orig", "new"):
                vals = data.SCORES[t][key]
                if len(vals) != len(PARAMS) or not all(1 <= v <= 5 for v in vals):
                    problems.append(f"turn {t}/{key}: bad score row {vals}")
            # Stiti's 12 Aug note: physiological grounding on the new responses
            # has to reach at least 3 - a 1 means the response ignored the body.
            physio = data.SCORES[t]["new"][PARAMS.index("PhysioGround")]
            if physio < 3:
                problems.append(f"turn {t}: new response PhysioGround is {physio}, minimum is 3")
    if problems:
        raise SystemExit("VALIDATION FAILED:\n  " + "\n  ".join(problems))
    print(f"  validated {len(turns)} turns")
    check_variation(data, turns)


# Channel keywords used to work out which channel each line of a summary reports,
# so the ordering of the five lines can be compared across turns.
CHANNEL_KEYS = (("stress", "Stress score"), ("eda", "EDA:"), ("pr", "PR:"),
                ("accel", "ACCEL:"))


def line_signature(summary):
    """Order the channels appear in, e.g. ('eda','accel','pr','stress').

    Channels now share lines, so the signature is the order of first mention
    across the whole summary rather than one channel per line.
    """
    found = []
    for name, marker in CHANNEL_KEYS:
        idx = summary.find(marker)
        if idx != -1:
            found.append((idx, name))
    return tuple(name for _, name in sorted(found))


def check_variation(data, turns):
    """Stiti's note: wording and sequence must differ from turn to turn.

    Repetition here is a real review finding (she spotted it on participant 9),
    so it is reported loudly rather than left for someone to notice downstream.
    """
    opens, sigs, windows = {}, {}, {}
    for t in turns:
        summary = data.PHYSIO_SHORT[t]
        first = summary.split("\n")[0]
        # The opening clause after "Physiological Signals:" is what reads as repetitive.
        body = first.split(":", 1)[1].strip() if ":" in first else first
        opens.setdefault(" ".join(body.split()[:4]).lower(), []).append(t)
        sigs.setdefault(line_signature(summary), []).append(t)
        # The window expression is whatever trails the EDA line after its last comma,
        # e.g. "over the past 6 minutes" / "first minute" / "this 6-minute stretch".
        for line in summary.split("\n"):
            if line.startswith("EDA:"):
                windows.setdefault(line.rsplit(",", 1)[-1].strip(" .<EOS>").lower(),
                                   []).append(t)
                break

    n = len(turns)
    worst_open = max(opens.items(), key=lambda kv: len(kv[1]))
    worst_sig = max(sigs.items(), key=lambda kv: len(kv[1]))
    worst_window = max(windows.items(), key=lambda kv: len(kv[1]))
    print(f"  variation: {len(opens)} distinct openings, {len(sigs)} line orders, "
          f"{len(windows)} window phrasings")

    warnings = []
    if len(opens) < max(3, n // 2):
        warnings.append(f"only {len(opens)} distinct openings across {n} turns "
                        f"(worst repeats {len(worst_open[1])}x: {worst_open[0]!r})")
    if len(sigs) < 3 or len(worst_sig[1]) > n * 0.5:
        warnings.append(f"line order repeats {len(worst_sig[1])}/{n} turns as {worst_sig[0]}")
    if len(worst_window[1]) > n * 0.6:
        warnings.append(f"time-window phrasing {worst_window[0]!r} used on "
                        f"{len(worst_window[1])}/{n} turns")
    for w in warnings:
        print(f"  WARNING repetition: {w}")


def md_cell(text):
    return text.replace("\n", "<br>").replace("|", "\\|")


def esc(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


CSS = """
:root{
  --ground:#F2F5F5; --surface:#FBFCFC; --ink:#10171A; --muted:#5B6B71;
  --line:#D6DFE0; --line-soft:#E6EDED; --accent:#0E6E75; --accent-soft:#DCEDED;
  --orig:#9C5322; --shadow:0 1px 2px rgba(16,23,26,.06);
}
@media (prefers-color-scheme: dark){
  :root:not([data-theme="light"]){
    --ground:#0E1417; --surface:#151D21; --ink:#E4EDEE; --muted:#8FA3A9;
    --line:#243036; --line-soft:#1D262B; --accent:#5AC3C6; --accent-soft:#12312F;
    --orig:#DB9160; --shadow:0 1px 2px rgba(0,0,0,.4);
  }
}
:root[data-theme="dark"]{
  --ground:#0E1417; --surface:#151D21; --ink:#E4EDEE; --muted:#8FA3A9;
  --line:#243036; --line-soft:#1D262B; --accent:#5AC3C6; --accent-soft:#12312F;
  --orig:#DB9160; --shadow:0 1px 2px rgba(0,0,0,.4);
}
*{box-sizing:border-box}
body{background:var(--ground); color:var(--ink);
  font-family:ui-sans-serif,-apple-system,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;
  font-size:16px; line-height:1.6; margin:0; padding:2.5rem 1.25rem 5rem}
.wrap{max-width:60rem; margin:0 auto; display:flex; flex-direction:column; gap:3rem}
h1,h2{font-family:Charter,"Bitstream Charter","Iowan Old Style",Palatino,Georgia,serif;
  text-wrap:balance; margin:0; line-height:1.2; font-weight:600}
h1{font-size:2rem; letter-spacing:-.01em}
h2{font-size:1.35rem}
p{margin:0}
.eyebrow{font-size:.72rem; letter-spacing:.14em; text-transform:uppercase;
  color:var(--muted); font-weight:600}
header{display:flex; flex-direction:column; gap:.85rem;
  border-bottom:1px solid var(--line); padding-bottom:1.5rem}
.meta{display:flex; flex-wrap:wrap; gap:.5rem 2rem; font-size:.85rem; color:var(--muted)}
.meta b{color:var(--ink); font-weight:600; font-variant-numeric:tabular-nums}
section{display:flex; flex-direction:column; gap:1.25rem}
.sec-head{display:flex; flex-direction:column; gap:.35rem;
  border-left:2px solid var(--accent); padding-left:.85rem}
.sec-head p{color:var(--muted); font-size:.9rem; max-width:46rem}
.turns{display:flex; flex-direction:column; gap:1rem}
.turn{background:var(--surface); border:1px solid var(--line); border-radius:3px;
  box-shadow:var(--shadow); display:grid; grid-template-columns:4.5rem 1fr}
.turn-no{border-right:1px solid var(--line-soft); padding:1.1rem .5rem;
  display:flex; flex-direction:column; align-items:center; gap:.2rem}
.turn-no span:first-child{font-size:.62rem; letter-spacing:.12em; text-transform:uppercase;
  color:var(--muted)}
.turn-no span:last-child{font-family:Charter,"Iowan Old Style",Georgia,serif;
  font-size:1.5rem; font-variant-numeric:tabular-nums; line-height:1}
.turn-body{padding:1.1rem 1.25rem; display:flex; flex-direction:column; gap:.9rem; min-width:0}
.user{font-family:Charter,"Iowan Old Style",Georgia,serif; font-size:1.02rem}
.user::before{content:"\\201C"} .user::after{content:"\\201D"}
.physio{font-family:ui-monospace,SFMono-Regular,Menlo,Consolas,monospace; font-size:.76rem;
  line-height:1.65; color:var(--muted); background:var(--ground);
  border:1px solid var(--line-soft); border-radius:2px; padding:.7rem .8rem;
  white-space:pre-wrap; overflow-x:auto}
.chip{display:inline-flex; align-self:flex-start; background:var(--accent-soft);
  color:var(--accent); border-radius:2px; font-size:.7rem; font-weight:700;
  letter-spacing:.08em; text-transform:uppercase; padding:.25rem .55rem}
.resp{font-size:.95rem}
.tablewrap{overflow-x:auto; border:1px solid var(--line); border-radius:3px;
  background:var(--surface); box-shadow:var(--shadow)}
table{border-collapse:collapse; width:100%; font-size:.85rem;
  font-variant-numeric:tabular-nums}
th,td{padding:.5rem .7rem; text-align:left; border-bottom:1px solid var(--line-soft);
  white-space:nowrap}
thead th{font-size:.68rem; letter-spacing:.09em; text-transform:uppercase;
  color:var(--muted); font-weight:700; border-bottom:1px solid var(--line)}
tbody tr.new td{border-bottom:1px solid var(--line)}
td.num{text-align:center; width:3.6rem}
td.turn-cell{color:var(--muted)}
.ver{display:inline-flex; align-items:center; gap:.45rem; font-weight:600}
.ver::before{content:""; width:.4rem; height:.9rem; border-radius:1px; background:var(--orig)}
tr.new .ver::before{background:var(--accent)}
.s1{background:transparent} .s2{background:color-mix(in srgb,var(--accent) 8%,transparent)}
.s3{background:color-mix(in srgb,var(--accent) 16%,transparent)}
.s4{background:color-mix(in srgb,var(--accent) 26%,transparent)}
.s5{background:color-mix(in srgb,var(--accent) 38%,transparent)}
tfoot td{font-weight:700; border-top:1px solid var(--line); background:var(--ground)}
.notes{font-size:.88rem; color:var(--muted)}
.notes ul{margin:0; padding-left:1.1rem}
.notes li{margin-bottom:.4rem}
.notes code{font-family:ui-monospace,Menlo,monospace; font-size:.82em; color:var(--ink)}
footer{border-top:1px solid var(--line); padding-top:1rem; font-size:.8rem; color:var(--muted)}
@media (max-width:34rem){
  .turn{grid-template-columns:1fr}
  .turn-no{flex-direction:row; justify-content:flex-start; gap:.5rem; border-right:0;
    border-bottom:1px solid var(--line-soft); padding:.6rem 1.25rem}
  .turn-no span:last-child{font-size:1.1rem}
}
"""


def notes(data, turns):
    base = [
        "Every original response scores 1 on physiological grounding: the <code>ai_text</code> "
        "field never references bodily state. New responses also score 1 where the physiology was "
        "unremarkable and the master prompt's rule that physiology must stay subtle argued "
        "against forcing a mention.",
        "Only turns with a populated <code>stress_level</code> carry a Mild/Moderate/High label. "
        "Other turns report the integer score with the label marked not reported, per the "
        "no-fabrication rule.",
        f"Activity level thresholds: low below {LOW_MAX}, moderate {LOW_MAX} to {MODERATE_MAX}, "
        f"high above {MODERATE_MAX}, calibrated against the worked examples in the brief.",
        "Shortening used only the timestamp, raw_physio channels and stress score for each turn, "
        "as requested - not the whole JSON.",
    ]
    return base + list(getattr(data, "NOTES", []))


def write_reports(pid, data, outdir, turns, turns_in, original, avg):
    label = getattr(data, "LABEL", pid)
    sessions = sorted({turns_in[t]["session_id"] for t in turns})

    lines = [
        f"# {label}",
        "",
        f"Source: `{data.SOURCE}` - {len(sessions)} session(s), {len(turns)} turns. "
        "The JSON is unmodified (step 5 is pending verification).",
        "",
        "## Table 1 - Shortened physiological signals and new AI responses",
        "",
        "| Turn | User text | Shortened Physiological Signals | Strategy | New AI Response |",
        "| --- | --- | --- | --- | --- |",
    ]
    for t in turns:
        strat, resp = data.NEW_RESPONSE[t]
        lines.append(f"| {t} | {md_cell(turns_in[t]['user_text'])} | "
                     f"{md_cell(data.PHYSIO_SHORT[t])} | {strat} | {md_cell(resp)} |")
    lines += [
        "",
        "## Table 2 - Evaluation scores (1-5, per turn)",
        "",
        "| Turn | Version | " + " | ".join(PARAMS) + " |",
        "| --- | --- | " + " | ".join(["---"] * len(PARAMS)) + " |",
    ]
    for t in turns:
        for v, key in (("Original AI Response", "orig"), ("New AI Response", "new")):
            lines.append(f"| {t} | {v} | " + " | ".join(str(x) for x in data.SCORES[t][key]) + " |")
    for v in avg:
        lines.append(f"| **Avg** | **{v}** | " + " | ".join(f"**{x}**" for x in avg[v]) + " |")
    lines += ["", "## Notes", ""]
    lines += [f"- {n.replace('<code>', '`').replace('</code>', '`')}" for n in notes(data, turns)]
    (outdir / f"{pid}_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    cards = "".join(
        f'<article class="turn"><div class="turn-no"><span>Turn</span><span>{t}</span></div>'
        f'<div class="turn-body"><p class="user">{esc(turns_in[t]["user_text"])}</p>'
        f'<pre class="physio">{esc(data.PHYSIO_SHORT[t])}</pre>'
        f'<span class="chip">{esc(data.NEW_RESPONSE[t][0])}</span>'
        f'<p class="resp">{esc(data.NEW_RESPONSE[t][1])}</p></div></article>'
        for t in turns)

    trs = []
    for t in turns:
        for v, key, cls in (("Original AI Response", "orig", "orig"),
                            ("New AI Response", "new", "new")):
            cells = "".join(f'<td class="num s{s}">{s}</td>' for s in data.SCORES[t][key])
            first = f'<td class="turn-cell" rowspan="2">{t}</td>' if key == "orig" else ""
            trs.append(f'<tr class="{cls}">{first}<td><span class="ver">{v}</span></td>{cells}</tr>')
    foot = "".join(
        f'<tr><td></td><td><span class="ver">{v} average</span></td>'
        + "".join(f'<td class="num">{x}</td>' for x in avg[v]) + "</tr>" for v in avg)

    html = f"""<title>{esc(label)} - physio-grounded response evaluation</title>
<style>{CSS}</style>
<div class="wrap">
<header>
  <p class="eyebrow">Physiological signal x LLM response study</p>
  <h1>{esc(label)}</h1>
  <div class="meta">
    <span>Sessions <b>{esc(', '.join(sessions))}</b></span>
    <span>Turns <b>{len(turns)}</b></span>
    <span>Source <b>{esc(data.SOURCE)}</b></span>
    <span>JSON <b>unmodified</b></span>
  </div>
</header>
<section>
  <div class="sec-head">
    <h2>Shortened physiological signals and new AI responses</h2>
    <p>Physiology re-summarised from the raw channels to five lines per turn; responses
    regenerated from the user text plus the full physiological summary using the master prompt,
    one counselling strategy per turn.</p>
  </div>
  <div class="turns">{cards}</div>
</section>
<section>
  <div class="sec-head">
    <h2>Evaluation scores</h2>
    <p>Original response from the <code>ai_text</code> field versus the newly generated response,
    scored 1 to 5 on each parameter, per turn.</p>
  </div>
  <div class="tablewrap"><table>
    <thead><tr><th>Turn</th><th>Version</th>{''.join(f'<th>{p}</th>' for p in PARAMS)}</tr></thead>
    <tbody>{''.join(trs)}</tbody>
    <tfoot>{foot}</tfoot>
  </table></div>
</section>
<section class="notes">
  <div class="sec-head"><h2>Notes on the data</h2></div>
  <ul>{''.join(f'<li>{n}</li>' for n in notes(data, turns))}</ul>
</section>
<footer>Step 5 (writing Physiological Signals, Strategy and Response back into
<code>training_string_case2</code>) is pending verification of these scores.</footer>
</div>
"""
    (outdir / f"{pid}_report.html").write_text(html, encoding="utf-8")


def write_docx(pid, data, outdir, turns, turns_in, avg):
    """Word file matching the house format of the shared Google Doc.

    That format is this generator's output after a Google Docs round-trip:
    Avenir throughout, plain "Participant N" title, a comma-separated meta line,
    both tables running on without a page break, and no notes section in the
    document itself (the caveats live in the .md and .html reports instead).
    """
    label = getattr(data, "LABEL", pid)
    sessions = sorted({turns_in[t]["session_id"] for t in turns})
    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin = sec.bottom_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)

    doc.add_heading(label, level=0).runs[0].font.name = FONT
    session_word = "Session" if len(sessions) == 1 else "Sessions"
    meta = doc.add_paragraph().add_run(
        f"{session_word} {', '.join(sessions)},  {len(turns)} turns, Source: {data.SOURCE}")
    meta.font.name = FONT

    def add_table(headers, widths):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.autofit = False
        for cell, head, w in zip(t.rows[0].cells, headers, widths):
            cell.width = Inches(w)
            run = cell.paragraphs[0].add_run(head)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = FONT
        return t

    def put(cell, text, width, size=9, mono=False, bold=False, center=False):
        cell.width = Inches(width)
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, line in enumerate(text.split("\n")):
            if i:
                p = cell.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(size)
            run.bold = bold
            run.font.name = FONT
            # The physiological summaries stay grey to separate them from prose.
            if mono:
                run.font.color.rgb = RGBColor(0x44, 0x4C, 0x50)

    def heading(text, center=False):
        h = doc.add_heading(text, level=1)
        h.runs[0].font.name = FONT
        if center:
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return h

    heading("Table 1 - Shortened physiological signals and new AI responses")
    w1 = [0.45, 2.1, 3.5, 1.35, 3.6]
    t1 = add_table(["Turn", "User text", "Shortened Physiological Signals",
                    "Strategy", "New AI Response"], w1)
    for turn in turns:
        strategy, response = data.NEW_RESPONSE[turn]
        cells = t1.add_row().cells
        put(cells[0], str(turn), w1[0], center=True, bold=True)
        put(cells[1], turns_in[turn]["user_text"], w1[1])
        put(cells[2], data.PHYSIO_SHORT[turn], w1[2], size=8, mono=True)
        put(cells[3], strategy, w1[3], bold=True)
        put(cells[4], response, w1[4])

    doc.add_paragraph()
    heading("Table 2 - Evaluation scores (1 to 5, per turn)", center=True)
    w2 = [0.6, 2.2, *[1.0] * len(PARAMS)]
    t2 = add_table(["Turn", "Version", *PARAMS], w2)
    for turn in turns:
        for version, key in (("Original AI Response", "orig"), ("New AI Response", "new")):
            cells = t2.add_row().cells
            put(cells[0], str(turn) if key == "orig" else "", w2[0], center=True)
            put(cells[1], version, w2[1])
            for cell, score in zip(cells[2:], data.SCORES[turn][key]):
                put(cell, str(score), 1.0, center=True)
    for version in avg:
        cells = t2.add_row().cells
        put(cells[0], "", w2[0])
        put(cells[1], f"{version} - average", w2[1], bold=True)
        for cell, value in zip(cells[2:], avg[version]):
            put(cell, str(value), 1.0, center=True, bold=True)

    doc.save(outdir / f"{pid}_report.docx")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("participants", nargs="+", help="e.g. P11 P11_1 P14")
    ap.add_argument("-o", "--outdir", default=str(ROOT / "out"))
    args = ap.parse_args()

    for pid in args.participants:
        outdir = Path(args.outdir) / pid
        data = load_data(pid)
        turns_in = read_csv(outdir, "turns_input.csv")
        original = read_csv(outdir, "original_ai.csv")
        turns = sorted(turns_in)
        print(f"{pid}: {len(turns)} turns")
        check(data, turns)

        write_csv(outdir, "physio_short.csv", ["turn", "shortened_physio_signals"],
                  [[t, data.PHYSIO_SHORT[t]] for t in turns])
        write_csv(outdir, "new_responses.csv", ["turn", "strategy", "response"],
                  [[t, *data.NEW_RESPONSE[t]] for t in turns])
        write_csv(outdir, "comparison.csv",
                  ["turn", "user_text", "original_ai_response", "new_ai_response"],
                  [[t, turns_in[t]["user_text"], original[t]["ai_text"],
                    f"Strategy: {data.NEW_RESPONSE[t][0]}\n\nResponse: {data.NEW_RESPONSE[t][1]}"]
                   for t in turns])
        write_csv(outdir, "evaluation.csv", ["turn", "version", *PARAMS],
                  [row for t in turns for row in (
                      [t, "Original AI Response", *data.SCORES[t]["orig"]],
                      [t, "New AI Response", *data.SCORES[t]["new"]])])

        avg = {}
        for v, key in (("Original AI Response", "orig"), ("New AI Response", "new")):
            avg[v] = [round(sum(data.SCORES[t][key][i] for t in turns) / len(turns), 2)
                      for i in range(len(PARAMS))]

        write_reports(pid, data, outdir, turns, turns_in, original, avg)
        write_docx(pid, data, outdir, turns, turns_in, avg)
        print(f"  wrote {pid}_report.md / .html / .docx")
        print("  avg orig:", dict(zip(PARAMS, avg["Original AI Response"])))
        print("  avg new :", dict(zip(PARAMS, avg["New AI Response"])))


if __name__ == "__main__":
    main()
